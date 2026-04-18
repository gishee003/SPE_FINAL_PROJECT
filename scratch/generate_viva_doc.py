from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

add_heading(doc, "SPE Final Project Viva Preparation Guide", 0)
add_paragraph(
    doc,
    "Project: Drift-Aware Churn Prediction Pipeline | Stack: Flask, Docker, Kubernetes, Jenkins, Ansible, ELK",
)

add_heading(doc, "1. Project Overview", 1)
add_paragraph(
    doc,
    "This project is an end-to-end MLOps-style microservices platform for bank customer churn prediction with drift-aware retraining. "
    "The key idea is: data is ingested, predictions are served, drift is detected, and retraining is triggered when drift appears.",
)
add_bullet(doc, "Business goal: maintain prediction quality even when live data distribution changes over time.")
add_bullet(doc, "Engineering goal: production-style deployment, observability, CI/CD, and resilient Kubernetes behavior.")
add_bullet(doc, "Academic value: demonstrates software platform engineering principles applied to ML systems.")

add_heading(doc, "2. Why This Architecture Was Used", 1)
add_paragraph(doc, "Microservices were chosen to separate responsibilities and improve deployability.")
add_bullet(doc, "Data Ingestion service: validates schema and orchestrates downstream calls.")
add_bullet(doc, "Model Serving service: isolated prediction endpoint, easy to scale independently.")
add_bullet(doc, "Drift Detection service: independent statistical checks; can trigger retraining.")
add_bullet(doc, "Model Training component: run as a Kubernetes Job for batch, run-to-completion behavior.")
add_paragraph(
    doc,
    "Why not monolith? With microservices, each component scales and evolves independently, and failures are isolated.",
)

add_heading(doc, "3. End-to-End Data Flow (Runtime Behavior)", 1)
add_bullet(doc, "Client sends records to POST /ingest.")
add_bullet(doc, "Ingestion validates schema and forwards data to prediction and drift services.")
add_bullet(doc, "Serving returns predictions and probabilities from model artifacts on PVC.")
add_bullet(doc, "Drift service compares live batch with reference distribution.")
add_bullet(doc, "If drift_detected = true, drift creates a Kubernetes training Job.")
add_bullet(doc, "Training Job retrains and overwrites churn_model.pkl and reference_distribution.pkl on shared PVC.")
add_bullet(doc, "Subsequent predictions use refreshed model artifacts.")

add_heading(doc, "4. Core Technical Choices and Justifications", 1)
add_heading(doc, "4.1 Flask APIs", 2)
add_paragraph(
    doc,
    "Flask is lightweight and fast to build REST services. It is ideal for educational and prototype-to-production APIs.",
)
add_heading(doc, "4.2 Scikit-learn Pipeline", 2)
add_paragraph(
    doc,
    "Using a single Pipeline (preprocessing + logistic regression) prevents training-serving skew and keeps inference logic consistent.",
)
add_heading(doc, "4.3 Shared PV/PVC", 2)
add_paragraph(
    doc,
    "A shared volume stores model artifacts and reference statistics so serving/drift can access latest outputs from training.",
)
add_heading(doc, "4.4 HPA on API Services (Not Training)", 2)
add_paragraph(
    doc,
    "HPA is useful for stateless/online services (ingestion, serving, drift). Training is batch and now run as Jobs; HPA is intentionally avoided there.",
)
add_heading(doc, "4.5 Training as Kubernetes Job", 2)
add_paragraph(
    doc,
    "Jobs are designed for finite tasks. They provide retry semantics, completion status, and avoid always-on training pods.",
)
add_heading(doc, "4.6 ELK for Observability", 2)
add_paragraph(
    doc,
    "Structured JSON logs are shipped with Filebeat to Elasticsearch and visualized in Kibana for reliability monitoring and debugging.",
)
add_heading(doc, "4.7 Jenkins CI/CD + Ansible", 2)
add_paragraph(
    doc,
    "Jenkins automates tests, builds, pushes, and deploys; Ansible handles environment provisioning and secret management through Vault.",
)

add_heading(doc, "5. Kubernetes Behavior and Deployment Details", 1)
add_bullet(doc, "Deployments manage always-on APIs (ingestion, serving, drift).")
add_bullet(doc, "Services expose stable network identities for inter-service communication.")
add_bullet(doc, "Readiness probes gate traffic until API is healthy.")
add_bullet(doc, "Liveness probes restart unhealthy pods to self-heal.")
add_bullet(doc, "Rolling update strategy avoids full downtime during image upgrades.")
add_bullet(doc, "HPA scales selected services based on CPU/memory utilization thresholds.")
add_bullet(doc, "Training retraining path uses Job creation from drift service (RBAC-enabled).")

add_heading(doc, "6. Why LoadBalancer Was Not Mandatory", 1)
add_paragraph(
    doc,
    "In this project, internal service-to-service communication occurs inside the cluster. "
    "A LoadBalancer service is only required when exposing endpoints externally at cloud-provider level.",
)
add_bullet(doc, "ClusterIP/NodePort can be enough for local labs and controlled environments.")
add_bullet(doc, "Ingress or one LB at gateway level is usually preferred over LB per microservice.")

add_heading(doc, "7. Drift Detection Logic (Conceptual)", 1)
add_bullet(doc, "Numerical feature drift is approximated with KS-style tests.")
add_bullet(doc, "Label drift compares Exited distribution changes versus baseline.")
add_bullet(doc, "Drift summary includes per-feature evidence and overall drift_detected flag.")
add_bullet(doc, "On detected drift, retraining is initiated (Job mode in Kubernetes).")

add_heading(doc, "8. CI/CD Pipeline Behavior", 1)
add_bullet(doc, "Unit tests run for all services.")
add_bullet(doc, "Integration flow trains baseline then validates ingest/predict/drift pipeline.")
add_bullet(doc, "Docker images are built and pushed to registry.")
add_bullet(doc, "Kubernetes manifests applied: PV, PVC, RBAC, Deployments, Services, HPA.")
add_bullet(doc, "ELK setup Job provisions dashboards/index patterns.")

add_heading(doc, "9. Security and Reliability Considerations", 1)
add_bullet(doc, "Ansible Vault stores sensitive values securely.")
add_bullet(doc, "RBAC limits drift service permissions to only required Job actions.")
add_bullet(doc, "Readiness/liveness probes improve resilience.")
add_bullet(doc, "Centralized logs improve incident analysis and auditability.")
add_bullet(doc, "Controlled retraining path reduces accidental duplicate training runs.")

add_heading(doc, "10. Viva-Focused Trade-offs to Mention", 1)
add_bullet(doc, "Using PVC is simple but can be a bottleneck vs model registry/object storage.")
add_bullet(doc, "KS-based drift checks are practical but not full concept-drift guarantees.")
add_bullet(doc, "Job-based retraining is robust, but monitoring/tracing each run is important.")
add_bullet(doc, "Current system favors clarity and educational value over cloud-native complexity.")
add_bullet(doc, "Future scope: model versioning, canary rollout, asynchronous queue for retraining.")

add_heading(doc, "11. Important Definitions (Quick Glossary)", 1)
definitions = {
    "Microservice": "A small, independently deployable service focused on one business capability.",
    "API": "Application Programming Interface; defined HTTP endpoints for communication.",
    "Flask": "A lightweight Python web framework used to build REST services.",
    "Docker": "Container platform that packages app + dependencies for consistent execution.",
    "Kubernetes": "Container orchestration system for deployment, scaling, and management.",
    "Pod": "Smallest deployable unit in Kubernetes (one or more containers).",
    "Deployment": "Kubernetes controller for managing stateless long-running pods.",
    "Service": "Stable network endpoint for accessing pods.",
    "ClusterIP": "Service type accessible only inside cluster.",
    "NodePort": "Service type exposing app on each node IP + static port.",
    "LoadBalancer": "Service type provisioning external load balancer (typically cloud-managed).",
    "Ingress": "HTTP/HTTPS routing layer into cluster with host/path rules.",
    "PersistentVolume (PV)": "Cluster storage resource provisioned independently of pods.",
    "PersistentVolumeClaim (PVC)": "Request for storage by a pod.",
    "HPA": "Horizontal Pod Autoscaler; adjusts replica count based on metrics.",
    "Readiness Probe": "Health check deciding if pod can receive traffic.",
    "Liveness Probe": "Health check deciding if pod should be restarted.",
    "Job": "Kubernetes resource for run-to-completion tasks.",
    "CronJob": "Scheduled Kubernetes Job.",
    "RBAC": "Role-based access control for Kubernetes API permissions.",
    "CI/CD": "Continuous Integration/Continuous Delivery automation pipeline.",
    "Jenkins": "Automation server used for CI/CD execution.",
    "Ansible": "Configuration and provisioning automation tool.",
    "Ansible Vault": "Encryption mechanism for storing sensitive Ansible data.",
    "ELK Stack": "Elasticsearch, Logstash/Filebeat, Kibana for log analytics.",
    "Data Drift": "Change in input data distribution compared with training data.",
    "Label Drift": "Change in target label distribution over time.",
    "Concept Drift": "Change in relationship between features and target.",
    "Model Artifact": "Persisted model files used for inference.",
    "Training-Serving Skew": "Mismatch between train-time and serve-time preprocessing/logic.",
}
for term, definition in definitions.items():
    add_paragraph(doc, f"{term}: {definition}")

add_heading(doc, "12. Likely Viva Questions You Can Prepare For", 1)
add_bullet(doc, "Why did you move training from Deployment to Job?")
add_bullet(doc, "Why HPA is used for some services and not all?")
add_bullet(doc, "How do you ensure model consistency after retraining?")
add_bullet(doc, "What happens if drift is detected repeatedly?")
add_bullet(doc, "What are limitations of your drift detection method?")
add_bullet(doc, "How would you productionize further on cloud?")

add_heading(doc, "13. Suggested Revision Strategy", 1)
add_bullet(doc, "First pass: understand architecture and flow (Sections 1 to 4).")
add_bullet(doc, "Second pass: focus on Kubernetes behavior and CI/CD (Sections 5 to 9).")
add_bullet(doc, "Third pass: memorize glossary and trade-offs (Sections 10 and 11).")
add_bullet(doc, "Final pass: rehearse likely questions with concise 30 to 60 second answers.")

output_path = "SPE_Viva_Preparation_Guide.docx"
doc.save(output_path)
print(output_path)
